#!/usr/bin/env python3
"""
Generate documentation for Spring Boot RealWorld Example App
in the style of the Devin.ai customer case study pages.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a horizontal line to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run()
    run.add_break()

def create_documentation():
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(36)
    title_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    title_style.font.bold = True
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(24)
    h1_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    h1_style.font.bold = True
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(18)
    h2_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    h2_style.font.bold = True
    
    # Normal style
    normal_style = styles['Normal']
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # =========================================================================
    # TITLE SECTION
    # =========================================================================
    title = doc.add_heading('Spring Boot RealWorld Example App', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('A Full-Stack Application Built with Spring Boot + MyBatis')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    subtitle_run.font.italic = True
    
    doc.add_paragraph()
    
    # =========================================================================
    # KEY METRICS SECTION
    # =========================================================================
    metrics_table = doc.add_table(rows=1, cols=5)
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    metrics = [
        ('REST + GraphQL', 'Dual API support'),
        ('DDD', 'Domain-Driven Design'),
        ('CQRS', 'Read/Write separation'),
        ('JWT', 'Secure authentication'),
        ('SQLite', 'Lightweight database'),
    ]
    
    row = metrics_table.rows[0]
    for i, (metric, description) in enumerate(metrics):
        cell = row.cells[i]
        # Metric value
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(metric)
        run1.font.size = Pt(20)
        run1.font.bold = True
        run1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        # Description
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(description)
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # =========================================================================
    # ABOUT THE PROJECT
    # =========================================================================
    doc.add_heading('About the Project', level=1)
    
    about_text = """This codebase was created to demonstrate a fully fledged full-stack application built with Spring Boot + MyBatis including CRUD operations, authentication, routing, pagination, and more. It adheres to the RealWorld spec and API, making it compatible with various frontend implementations.

The application serves as a reference implementation for building production-ready applications using modern Java development practices, including Domain-Driven Design (DDD), Command Query Responsibility Segregation (CQRS), and both REST and GraphQL APIs."""
    
    about_para = doc.add_paragraph(about_text)
    about_para.paragraph_format.space_after = Pt(12)
    
    # Project info box
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Technology Stack', 'Spring Boot, MyBatis, GraphQL (Netflix DGS)'),
        ('Java Version', 'Java 11'),
        ('Repository', 'github.com/gothinkster/spring-boot-realworld-example-app'),
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        label_cell = row.cells[0]
        value_cell = row.cells[1]
        
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.bold = True
        label_run.font.size = Pt(10)
        
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(value)
        value_run.font.size = Pt(10)
        
        set_cell_shading(label_cell, 'F5F5F5')
    
    doc.add_paragraph()
    
    # =========================================================================
    # OVERVIEW
    # =========================================================================
    doc.add_heading('Overview', level=1)
    
    overview_text = """The Spring Boot RealWorld Example App demonstrates how to build a medium.com clone (called "Conduit") using Spring Boot and MyBatis. It implements the RealWorld API specification, which defines a standard set of endpoints for a blogging platform including user authentication, article management, comments, favorites, and user profiles.

What makes this implementation unique is its support for both REST and GraphQL APIs, allowing developers to choose their preferred approach or use both simultaneously. The codebase follows Domain-Driven Design principles to separate business logic from infrastructure concerns, making it highly maintainable and testable."""
    
    doc.add_paragraph(overview_text)
    
    # Quote box
    quote_para = doc.add_paragraph()
    quote_para.paragraph_format.left_indent = Inches(0.5)
    quote_para.paragraph_format.right_indent = Inches(0.5)
    quote_run = quote_para.add_run('"Following some DDD principles, REST or GraphQL is just a kind of adapter. And the domain layer will be consistent all the time."')
    quote_run.font.italic = True
    quote_run.font.size = Pt(12)
    quote_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    quote_attr = doc.add_paragraph()
    quote_attr.paragraph_format.left_indent = Inches(0.5)
    quote_attr_run = quote_attr.add_run('- Project Documentation')
    quote_attr_run.font.size = Pt(10)
    quote_attr_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    
    # =========================================================================
    # KEY FEATURES & ARCHITECTURE
    # =========================================================================
    doc.add_heading('Key Features & Architecture', level=1)
    
    features_table = doc.add_table(rows=5, cols=2)
    features_table.style = 'Table Grid'
    
    # Header row
    header_row = features_table.rows[0]
    header_row.cells[0].paragraphs[0].add_run('Component').bold = True
    header_row.cells[1].paragraphs[0].add_run('Description').bold = True
    set_cell_shading(header_row.cells[0], 'E8E8E8')
    set_cell_shading(header_row.cells[1], 'E8E8E8')
    
    features_data = [
        ('API Layer', 'REST endpoints via Spring MVC and GraphQL via Netflix DGS framework. Supports full CRUD operations for articles, comments, users, and profiles.'),
        ('Core Domain', 'Business entities (Article, User, Comment, Tag) and domain services. Implements the heart of the application logic independent of delivery mechanism.'),
        ('Application Services', 'Query services implementing CQRS pattern. Separates read operations from write operations for better scalability and maintainability.'),
        ('Infrastructure', 'MyBatis mappers for data persistence, JWT-based authentication with Spring Security, and SQLite database for easy local development.'),
    ]
    
    for i, (component, description) in enumerate(features_data):
        row = features_table.rows[i + 1]
        comp_cell = row.cells[0]
        desc_cell = row.cells[1]
        
        comp_para = comp_cell.paragraphs[0]
        comp_run = comp_para.add_run(component)
        comp_run.font.bold = True
        comp_run.font.size = Pt(10)
        
        desc_para = desc_cell.paragraphs[0]
        desc_run = desc_para.add_run(description)
        desc_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # =========================================================================
    # DEEP DIVE: CODE ORGANIZATION
    # =========================================================================
    doc.add_heading('Deep Dive: Code Organization', level=1)
    
    org_intro = """The codebase follows a clean architecture approach with clear separation of concerns:"""
    doc.add_paragraph(org_intro)
    
    # Code structure
    code_structure = doc.add_paragraph()
    code_structure.paragraph_format.left_indent = Inches(0.3)
    
    structure_items = [
        ('api/', 'Web layer implemented by Spring MVC - REST controllers and request/response handling'),
        ('core/', 'Business model including entities (Article, User, Comment) and domain services'),
        ('application/', 'High-level services for querying data transfer objects, implementing CQRS read models'),
        ('infrastructure/', 'Implementation classes for technical details - MyBatis mappers, repositories, JWT service'),
        ('graphql/', 'GraphQL datafetchers and mutations using Netflix DGS framework'),
    ]
    
    for folder, desc in structure_items:
        item_para = doc.add_paragraph()
        item_para.paragraph_format.left_indent = Inches(0.3)
        folder_run = item_para.add_run(folder)
        folder_run.font.bold = True
        folder_run.font.name = 'Courier New'
        folder_run.font.size = Pt(10)
        desc_run = item_para.add_run(f' - {desc}')
        desc_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # =========================================================================
    # DEEP DIVE: REST API
    # =========================================================================
    doc.add_heading('REST API Endpoints', level=1)
    
    api_intro = """The application exposes a comprehensive REST API following the RealWorld specification:"""
    doc.add_paragraph(api_intro)
    
    api_table = doc.add_table(rows=9, cols=3)
    api_table.style = 'Table Grid'
    
    # Header
    api_header = api_table.rows[0]
    api_header.cells[0].paragraphs[0].add_run('Endpoint').bold = True
    api_header.cells[1].paragraphs[0].add_run('Method').bold = True
    api_header.cells[2].paragraphs[0].add_run('Description').bold = True
    for cell in api_header.cells:
        set_cell_shading(cell, 'E8E8E8')
    
    api_endpoints = [
        ('/users', 'POST', 'Register a new user'),
        ('/users/login', 'POST', 'Authenticate user and get JWT token'),
        ('/user', 'GET/PUT', 'Get or update current user'),
        ('/profiles/:username', 'GET', 'Get user profile'),
        ('/articles', 'GET/POST', 'List or create articles'),
        ('/articles/:slug', 'GET/PUT/DELETE', 'Get, update, or delete article'),
        ('/articles/:slug/comments', 'GET/POST', 'List or add comments'),
        ('/tags', 'GET', 'Get all tags'),
    ]
    
    for i, (endpoint, method, desc) in enumerate(api_endpoints):
        row = api_table.rows[i + 1]
        
        ep_run = row.cells[0].paragraphs[0].add_run(endpoint)
        ep_run.font.name = 'Courier New'
        ep_run.font.size = Pt(9)
        
        method_run = row.cells[1].paragraphs[0].add_run(method)
        method_run.font.size = Pt(9)
        
        desc_run = row.cells[2].paragraphs[0].add_run(desc)
        desc_run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # =========================================================================
    # DEEP DIVE: GRAPHQL API
    # =========================================================================
    doc.add_heading('GraphQL API', level=1)
    
    graphql_intro = """In addition to REST, the application provides a full GraphQL API using Netflix's DGS framework. This allows clients to request exactly the data they need in a single query."""
    doc.add_paragraph(graphql_intro)
    
    doc.add_heading('Queries', level=2)
    queries_text = """The GraphQL schema supports the following queries: article (get single article by slug), articles (paginated list with filters), me (current user), feed (personalized article feed), profile (user profile by username), and tags (all available tags)."""
    doc.add_paragraph(queries_text)
    
    doc.add_heading('Mutations', level=2)
    mutations_text = """Available mutations include: createUser, login, updateUser, followUser, unfollowUser for user management; createArticle, updateArticle, deleteArticle, favoriteArticle, unfavoriteArticle for article operations; and addComment, deleteComment for comment management."""
    doc.add_paragraph(mutations_text)
    
    doc.add_paragraph()
    
    # =========================================================================
    # SECURITY
    # =========================================================================
    doc.add_heading('Security', level=1)
    
    security_text = """The application integrates with Spring Security and implements JWT (JSON Web Token) based authentication. When a user logs in or registers, they receive a JWT token that must be included in subsequent requests via the Authorization header.

Key security features include password encryption using BCrypt, stateless authentication via JWT tokens, configurable token expiration time, and protected endpoints requiring valid authentication."""
    doc.add_paragraph(security_text)
    
    doc.add_paragraph()
    
    # =========================================================================
    # DATABASE
    # =========================================================================
    doc.add_heading('Database', level=1)
    
    db_text = """The application uses SQLite as its default database for easy local development and testing. The database schema is managed through Flyway migrations, ensuring consistent database state across environments.

The data access layer is implemented using MyBatis with the Data Mapper pattern, providing clean separation between domain objects and database operations. The configuration can be easily changed in application.properties to use other databases like PostgreSQL or MySQL for production deployments."""
    doc.add_paragraph(db_text)
    
    doc.add_paragraph()
    
    # =========================================================================
    # GETTING STARTED
    # =========================================================================
    doc.add_heading('Getting Started', level=1)
    
    doc.add_heading('Prerequisites', level=2)
    prereq_para = doc.add_paragraph()
    prereq_para.add_run('Java 11 or higher must be installed on your system.')
    
    doc.add_heading('Running the Application', level=2)
    run_para = doc.add_paragraph()
    run_cmd = run_para.add_run('./gradlew bootRun')
    run_cmd.font.name = 'Courier New'
    run_cmd.font.size = Pt(10)
    
    run_desc = doc.add_paragraph('The application will start on http://localhost:8080. You can verify it is running by accessing http://localhost:8080/tags in your browser.')
    
    doc.add_heading('Running with Docker', level=2)
    docker_para = doc.add_paragraph()
    docker_cmd1 = docker_para.add_run('./gradlew bootBuildImage --imageName spring-boot-realworld-example-app')
    docker_cmd1.font.name = 'Courier New'
    docker_cmd1.font.size = Pt(10)
    
    docker_para2 = doc.add_paragraph()
    docker_cmd2 = docker_para2.add_run('docker run -p 8081:8080 spring-boot-realworld-example-app')
    docker_cmd2.font.name = 'Courier New'
    docker_cmd2.font.size = Pt(10)
    
    doc.add_heading('Running Tests', level=2)
    test_para = doc.add_paragraph()
    test_cmd = test_para.add_run('./gradlew test')
    test_cmd.font.name = 'Courier New'
    test_cmd.font.size = Pt(10)
    
    test_desc = doc.add_paragraph('The repository contains comprehensive test cases covering both API tests and repository tests.')
    
    doc.add_heading('Code Formatting', level=2)
    format_para = doc.add_paragraph()
    format_cmd = format_para.add_run('./gradlew spotlessJavaApply')
    format_cmd.font.name = 'Courier New'
    format_cmd.font.size = Pt(10)
    
    format_desc = doc.add_paragraph('The project uses Spotless with Google Java Format for consistent code style.')
    
    doc.add_paragraph()
    
    # =========================================================================
    # CONTRIBUTING
    # =========================================================================
    doc.add_heading('Contributing', level=1)
    
    contrib_text = """Contributions are welcome! Please fork the repository and submit pull requests to improve the project. Whether it is fixing bugs, improving documentation, or adding new features, all contributions help make this reference implementation better for the community."""
    doc.add_paragraph(contrib_text)
    
    # Save the document
    doc.save('docs/Spring_Boot_RealWorld_Documentation.docx')
    print('Documentation generated successfully: docs/Spring_Boot_RealWorld_Documentation.docx')

if __name__ == '__main__':
    create_documentation()
